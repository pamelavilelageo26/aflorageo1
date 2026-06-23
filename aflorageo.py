import streamlit as st
import sqlite3
import os
import json
import datetime
import re
import math
import io
import base64
import zipfile
from pathlib import Path
import folium
from folium.plugins import MarkerCluster
from streamlit_folium import st_folium
import pandas as pd
import utm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# CONFIGURAÇÕES GLOBAIS
# ---------------------------------------------------------------------------
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "aflorageo.db")
DEFAULT_FREE_LIMIT = 30
ACTIVATION_KEY = "AFLORAGEO-PREMIUM-2024"

CORES_MUNSELL = [
    ("", "Selecione a cor"),
    ("N1", "Preto"),
    ("N2", "Cinza muito escuro"),
    ("N3", "Cinza escuro"),
    ("N4", "Cinza médio escuro"),
    ("N5", "Cinza médio"),
    ("N6", "Cinza médio claro"),
    ("N7", "Cinza claro"),
    ("N8", "Cinza muito claro"),
    ("N9", "Branco"),
    ("5R 3/2", "Vermelho escuro acinzentado"),
    ("5R 3/4", "Vermelho escuro"),
    ("5R 4/6", "Vermelho"),
    ("5R 5/6", "Vermelho claro"),
    ("10R 3/4", "Vermelho escuro"),
    ("10R 4/2", "Vermelho escuro acinzentado"),
    ("10R 4/6", "Vermelho médio"),
    ("10R 5/4", "Vermelho claro acinzentado"),
    ("5YR 3/4", "Marrom avermelhado escuro"),
    ("5YR 4/4", "Marrom avermelhado"),
    ("5YR 5/6", "Marrom claro avermelhado"),
    ("5YR 6/3", "Marrom claro"),
    ("5YR 7/4", "Marrom amarelado claro"),
    ("7.5YR 4/4", "Marrom"),
    ("7.5YR 5/6", "Marrom claro"),
    ("7.5YR 6/4", "Marrom claro acinzentado"),
    ("10YR 3/4", "Marrom amarelado escuro"),
    ("10YR 4/4", "Marrom amarelado"),
    ("10YR 5/4", "Marrom amarelado"),
    ("10YR 5/6", "Marrom amarelado"),
    ("10YR 6/3", "Marrom claro acinzentado"),
    ("10YR 7/4", "Amarelo claro acinzentado"),
    ("10YR 8/1", "Branco"),
    ("2.5Y 5/4", "Marrom oliva claro"),
    ("2.5Y 6/3", "Cinza claro amarelado"),
    ("2.5Y 7/2", "Cinza claro amarelado"),
    ("2.5Y 8/1", "Branco acinzentado"),
    ("5Y 4/1", "Cinza escuro"),
    ("5Y 5/2", "Cinza oliva claro"),
    ("5Y 6/2", "Cinza oliva claro"),
    ("5Y 7/3", "Amarelo claro"),
    ("5GY 5/2", "Verde acinzentado"),
    ("5G 5/3", "Verde acinzentado"),
    ("5BG 5/3", "Azul esverdeado acinzentado"),
    ("5B 5/4", "Azul acinzentado"),
    ("5PB 5/3", "Azul arroxeado acinzentado"),
    ("5P 5/4", "Roxo acinzentado"),
    ("5RP 5/3", "Roxo avermelhado acinzentado"),
    ("5R 6/2", "Rosa acinzentado"),
]

TIPOS_AFLO = [
    "Corte de estrada", "Lajeado", "Pedreira", "Rio", "Trincheira", "Galeria",
    "Afloramento natural", "Barranco", "Talude", "Túnel", "Poço", "Mina", "Açude",
    "Praia", "Dunas", "Terraço", "Voçoroca", "Bloco isolado (matacão)",
    "Cavidade artificial (empresta)", "Outro"
]

ACESSOS = [
    "Caminhamento", "Carro 4x4", "Carro passeio", "Barco motorizado", "Barco a remo",
    "Helicóptero", "Avião", "Bicicleta", "Moto", "Ônibus", "Animal (cavalo/burro/mula)",
    "Trator", "Quadriciclo"
]

TIPOS_ESTRUTURA_SELECT = [
    "Acamamento (Bedding)", "Foliação (Foliation)", "Fratura (Joint)", "Falha (Fault)",
    "Dobra (Fold)", "Lineação (Lineation)", "Paleocorrente (Paleocurrent)", "Veio (Vein)",
    "Dique (Dike)", "Xistosidade (Schistosity)", "Clivagem (Cleavage)", "Eixo de dobra (Fold axis)",
    "Superfície axial (Axial surface)", "Zona de cisalhamento (Shear zone)", "Contato (Contact)",
    "Outro (Other)"
]

FINALIDADES_AMOSTRA_LISTA = [
    "Petrografia", "Geoquímica (rocha total)", "Geoquímica (orgânica)", "Geocronologia",
    "Granulometria", "Difração de Raios-X (DRX)", "Fluorescência de Raios-X (FRX)",
    "Microscopia Eletrônica (MEV/EDS)", "Isótopos Estáveis", "Datação Radiométrica",
    "Paleontologia", "Análise Mineralógica", "Análise Petrofísica", "Microssonda Eletrônica",
    "Laminação Delgada", "Catodoluminescência", "Bioestratigrafia", "Sedimentologia",
    "Geofísica de Poço", "Análise Térmica (DSC/TGA)", "Outro"
]

DENSIDADE_ROCHAS = [
    ("Aluvião", "1,96-2,00", "1,98"),
    ("Argila", "1,63-2,60", "2,21"),
    ("Areia", "1,70-2,30", "2,0"),
    ("Arenito", "1,61-2,76", "2,35"),
    ("Argilito", "1,77-3,20", "2,4"),
    ("Calcário", "1,93-2,90", "2,55"),
    ("Riolito", "2,35-2,70", "2,52"),
    ("Andesito", "2,40-2,80", "2,61"),
    ("Granito", "2,50-2,81", "2,64"),
    ("Granodiorito", "2,67-2,79", "2,73"),
    ("Diabásio", "2,50-3,20", "2,91"),
    ("Basalto", "2,70-3,30", "2,99"),
    ("Gabro", "2,70-3,50", "3,03"),
    ("Peridotito", "2,78-3,37", "3,15"),
    ("Piroxenito", "2,93-3,34", "3,17"),
    ("Quartzito", "2,50-2,70", "2,6"),
    ("Xisto", "2,39-2,90", "2,64"),
    ("Granulito", "2,52-2,73", "2,65"),
    ("Filito", "2,68-2,80", "2,74"),
    ("Mármore", "2,60-2,90", "2,75"),
    ("Ardósia", "2,70-2,90", "2,78"),
    ("Gnaisse", "2,59-3,00", "2,8"),
    ("Anfibolito", "2,90-3,04", "2,96"),
    ("Eclogito", "3,20-3,54", "3,37"),
]

VELOCIDADE_ONDAS_P = [
    ("Areia (seca)", "0,2 - 1,0"),
    ("Areia (saturada em água)", "1,5 - 2,0"),
    ("Argila", "1,0 - 2,5"),
    ("Till glacial (saturado em água)", "1,5 - 2,5"),
    ("Permafroste", "3,5 - 4,0"),
    ("Arenitos", "2,0 - 6,0"),
    ("Arenito Terciário", "2,0 - 2,5"),
    ("Arenito Pennant (Carbonífero)", "4,0 - 4,5"),
    ("Quartzito Cambriano", "5,5 - 6,0"),
    ("Calcários", "2,0 - 6,0"),
    ("Greda Cretácea", "2,0 - 2,5"),
    ("Oólitos Jurássicos e Calcários Bioclásticos", "3,0 - 4,0"),
    ("Calcário Carbonífero", "5,0 - 5,5"),
    ("Dolomitos", "2,5 - 6,5"),
    ("Sal", "4,5 - 5,0"),
    ("Anidrita", "4,5 - 6,5"),
    ("Gipso", "2,0 - 3,5"),
    ("Granito", "5,5 - 6,0"),
    ("Gabro", "6,5 - 7,0"),
    ("Rochas ultramáficas", "7,5 - 8,5"),
    ("Serpentinito", "5,5 - 6,5"),
    ("Ar", "0,3"),
    ("Água", "1,4 - 1,5"),
    ("Gelo", "3,4"),
    ("Petróleo", "1,3 - 1,4"),
    ("Aço", "6,1"),
    ("Ferro", "5,8"),
    ("Alumínio", "6,6"),
    ("Concreto", "3,6"),
]

SUSCETIBILIDADE_MAGNETICA = [
    ("Ar", "~0"),
    ("Quartzo", "–0,01"),
    ("Rocha de sal", "–0,01"),
    ("Calcita", "–0,001 a 0,01"),
    ("Esfalerita", "–0,4"),
    ("Pirita", "–0,05 a 5"),
    ("Hematita", "0,5 a 35"),
    ("Ilmenita", "300 a 3500"),
    ("Magnetita", "1200 a 19200"),
    ("Calcário", "0 – 3"),
    ("Arenito", "0 – 20"),
    ("Folhelho", "0,01 – 15"),
    ("Xisto", "0,3 – 3"),
    ("Gnaisse", "0,1 – 25"),
    ("Ardósia", "0 – 35"),
    ("Granito", "0 – 50"),
    ("Gabro", "1 – 90"),
    ("Basalto", "0,2 – 175"),
    ("Peridotito", "90 – 200"),
]

CLASSES_DECLIVIDADE = [
    ("Plano", "0 – 3%", "Superfície horizontal"),
    ("Suave ondulado", "3 – 8%", "Suave inclinação"),
    ("Ondulado", "8 – 20%", "Relevo ondulado"),
    ("Forte ondulado", "20 – 45%", "Aclive acentuado"),
    ("Montanhoso", "45 – 75%", "Relevo escarpado"),
    ("Escarpado", "> 75%", "Paredões, penhascos"),
]

PADROES_DRENAGEM = [
    ("Dendrítico", "Ramificação aleatória", "Rochas homogêneas"),
    ("Paralelo", "Canais paralelos", "Declives acentuados"),
    ("Retangular", "Canais em ângulos retos", "Fraturas controlantes"),
    ("Trellis / Treliçado", "Canais paralelos com tributários", "Estruturas dobradas"),
    ("Radial", "Drenagem divergente", "Cúpulas, cones vulcânicos"),
    ("Centripeta", "Drenagem convergente", "Depressões"),
    ("Anelar", "Canais concêntricos", "Estruturas circulares"),
    ("Pinnate / Pinnado", "Drenagem em pena", "Maciços de granito"),
]

FORMAS_RELEVO = [
    ("Planície", "Apf", "Superfície plana"),
    ("Planalto", "Pl", "Superfície elevada com topo plano"),
    ("Depressão", "Dep", "Superfície rebaixada"),
    ("Cuesta", "Cu", "Escarpa com declive suave e escarpado"),
    ("Morro", "Mo", "Elevação arredondada"),
    ("Serra", "Se", "Elevação alongada com cume estreito"),
    ("Montanha", "Mt", "Elevação com grande altimetria"),
    ("Vale", "Va", "Depressão linear entre elevações"),
    ("Vale em V", "Vv", "Vale com encostas em V"),
    ("Vale em U", "Vu", "Vale com encostas em U"),
    ("Cumeada", "Cu", "Linha de cume alongada"),
    ("Escarpa", "Es", "Desnível abrupto"),
    ("Escarpa costeira", "Ecs", "Escarpa junto ao litoral"),
]

INTEMPERISMO_IBGE = [
    ("FR (Fresco)", "Rocha sã", "Sem alteração visível"),
    ("LE (Levemente alterado)", "Alteração incipiente", "Alteração superficial"),
    ("MO (Moderado)", "Alteração parcial", "Minerais alterados parcialmente"),
    ("AL (Altamente alterado)", "Alteração intensa", "Textura parcialmente preservada"),
    ("DE (Decomposto)", "Alteração total", "Textura praticamente destruída"),
    ("SA (Saprolito)", "Totalmente alterado", "Mantém relíquias de textura"),
    ("SO (Solo)", "Material de solo", "Sem textura rochosa preservada"),
    ("Outro", "Outro tipo de intemperismo", "Descrever nas observações"),
]


# ---------------------------------------------------------------------------
# UTILITÁRIO: RERUN COMPATÍVEL
# ---------------------------------------------------------------------------
def safe_rerun():
    """st.rerun() com fallback para versões antigas do Streamlit."""
    try:
        st.rerun()
    except AttributeError:
        st.experimental_rerun()


# ---------------------------------------------------------------------------
# CSS PERSONALIZADO
# ---------------------------------------------------------------------------
def inject_css():
    st.markdown(
        """
        <style>
        .main { background-color: #f8f9fa; }
        h1 { color: #2c3e50; }
        h2 { color: #34495e; }
        h3 { color: #34495e; }
        .stButton>button { border-radius: 6px; }
        .stProgress > div > div > div > div { background-color: #27ae60; }
        .premium-badge { color: #f39c12; font-weight: bold; }
        .free-badge { color: #3498db; font-weight: bold; }
        .block-container { padding-top: 2rem; padding-bottom: 2rem; }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ---------------------------------------------------------------------------
# BANCO DE DADOS
# ---------------------------------------------------------------------------
def get_connection():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def init_db():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS stations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ponto_id TEXT UNIQUE,
                data TEXT,
                utm_zone TEXT,
                hemisferio TEXT,
                utm_east REAL,
                utm_north REAL,
                latitude REAL,
                longitude REAL,
                altitude REAL,
                localizacao TEXT,
                municipio TEXT,
                contexto_geologico TEXT,
                tipo_afloramento TEXT,
                dimensoes TEXT,
                orientacao_afloramento TEXT,
                acesso TEXT,
                litologia_principal TEXT,
                litologia_secundaria TEXT,
                granulometria TEXT,
                cor TEXT,
                intemperismo TEXT,
                observacoes TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS structures (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                station_id INTEGER,
                tipo TEXT,
                strike TEXT,
                dip REAL,
                dip_dir REAL,
                plunge REAL,
                azimuth REAL,
                observacoes TEXT,
                created_at TEXT
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS samples (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                station_id INTEGER,
                codigo TEXT,
                tipo TEXT,
                finalidade TEXT,
                orientada INTEGER,
                observacoes TEXT,
                created_at TEXT
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS photos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                station_id INTEGER,
                descricao TEXT,
                arquivo TEXT,
                created_at TEXT
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS license (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                is_premium INTEGER DEFAULT 0,
                stations_limit INTEGER DEFAULT 30,
                user_email TEXT,
                plan_type TEXT,
                expires_at TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        """)
        conn.commit()
        conn.close()
        init_license()
    except Exception as e:
        st.error(f"Erro ao inicializar banco de dados: {e}")


def init_license():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT id FROM license LIMIT 1")
        if cur.fetchone() is None:
            now = datetime.datetime.now().isoformat()
            cur.execute(
                "INSERT INTO license (is_premium, stations_limit, user_email, plan_type, expires_at, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                (0, DEFAULT_FREE_LIMIT, "", "free", "", now, now),
            )
            conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao inicializar licença: {e}")


# ---------------------------------------------------------------------------
# LICENÇA
# ---------------------------------------------------------------------------
def get_license():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM license ORDER BY id DESC LIMIT 1")
        row = cur.fetchone()
        cols = [desc[0] for desc in cur.description] if cur.description else []
        lic = dict(zip(cols, row)) if row else None
        conn.close()
        return lic
    except Exception as e:
        st.error(f"Erro ao consultar licença: {e}")
        return None


def is_premium():
    lic = get_license()
    return bool(lic and lic.get("is_premium") == 1)


def get_station_count():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM stations")
        count = cur.fetchone()[0]
        conn.close()
        return count
    except Exception as e:
        st.error(f"Erro ao contar estações: {e}")
        return 0


def can_add_station():
    lic = get_license()
    if not lic:
        return False
    if lic.get("is_premium") == 1:
        return True
    limit = lic.get("stations_limit", DEFAULT_FREE_LIMIT)
    return get_station_count() < limit


def activate_premium(plan_type, email):
    try:
        now = datetime.datetime.now()
        if plan_type == "Mensal":
            expires = now + datetime.timedelta(days=30)
        elif plan_type == "Semestral":
            expires = now + datetime.timedelta(days=180)
        elif plan_type == "Anual":
            expires = now + datetime.timedelta(days=365)
        else:
            expires = now + datetime.timedelta(days=30)
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM license")
        cur.execute(
            "INSERT INTO license (is_premium, stations_limit, user_email, plan_type, expires_at, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (1, 999999, email, plan_type, expires.isoformat(), now.isoformat(), now.isoformat()),
        )
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Erro ao ativar premium: {e}")
        return False


def reset_license():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM license")
        now = datetime.datetime.now().isoformat()
        cur.execute(
            "INSERT INTO license (is_premium, stations_limit, user_email, plan_type, expires_at, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (0, DEFAULT_FREE_LIMIT, "", "free", "", now, now),
        )
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Erro ao redefinir licença: {e}")
        return False


# ---------------------------------------------------------------------------
# UTM / COORDENADAS
# ---------------------------------------------------------------------------
def utm_to_latlon(zone, east, north, hemisferio):
    try:
        zone_number = int(re.sub(r"[^0-9]", "", str(zone)) or "0")
        if zone_number < 1 or zone_number > 60:
            return None, None
        northern = (hemisferio and hemisferio.upper().startswith("N"))
        lat, lon = utm.to_latlon(east, north, zone_number, northern=northern)
        return lat, lon
    except Exception:
        return None, None


def latlon_to_utm(lat, lon):
    try:
        res = utm.from_latlon(lat, lon)
        return res[2], res[0], res[1]
    except Exception:
        return None, None, None


# ---------------------------------------------------------------------------
# IDs
# ---------------------------------------------------------------------------
def next_station_id():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT ponto_id FROM stations ORDER BY id DESC LIMIT 1")
        row = cur.fetchone()
        conn.close()
        if row and row[0]:
            m = re.search(r"(\d+)", str(row[0]))
            if m:
                n = int(m.group(1)) + 1
            else:
                n = get_station_count() + 1
        else:
            n = 1
        return f"AF-{n:03d}"
    except Exception as e:
        st.error(f"Erro ao gerar ID: {e}")
        return f"AF-{get_station_count() + 1:03d}"


# ---------------------------------------------------------------------------
# CRUD ESTAÇÃO
# ---------------------------------------------------------------------------
def insert_station(data):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        data["created_at"] = now
        data["updated_at"] = now
        cols = [
            "ponto_id", "data", "utm_zone", "hemisferio", "utm_east", "utm_north",
            "latitude", "longitude", "altitude", "localizacao", "municipio",
            "contexto_geologico", "tipo_afloramento", "dimensoes", "orientacao_afloramento",
            "acesso", "litologia_principal", "litologia_secundaria", "granulometria", "cor",
            "intemperismo", "observacoes", "created_at", "updated_at"
        ]
        values = [data.get(c) for c in cols]
        cur.execute(
            f"INSERT INTO stations ({','.join(cols)}) VALUES ({','.join(['?']*len(cols))})",
            values,
        )
        station_id = cur.lastrowid
        conn.commit()
        conn.close()
        return station_id
    except Exception as e:
        st.error(f"Erro ao salvar estação: {e}")
        return None


def update_station(station_id, data):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        data["updated_at"] = now
        cols = [
            "ponto_id", "data", "utm_zone", "hemisferio", "utm_east", "utm_north",
            "latitude", "longitude", "altitude", "localizacao", "municipio",
            "contexto_geologico", "tipo_afloramento", "dimensoes", "orientacao_afloramento",
            "acesso", "litologia_principal", "litologia_secundaria", "granulometria", "cor",
            "intemperismo", "observacoes", "updated_at"
        ]
        set_clause = ",".join([f"{c}=?" for c in cols])
        values = [data.get(c) for c in cols] + [station_id]
        cur.execute(f"UPDATE stations SET {set_clause} WHERE id=?", values)
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Erro ao atualizar estação: {e}")
        return False


def delete_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        for t in ["structures", "samples", "photos"]:
            cur.execute(f"DELETE FROM {t} WHERE station_id=?", (station_id,))
        cur.execute("DELETE FROM stations WHERE id=?", (station_id,))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Erro ao excluir estação: {e}")
        return False


def get_station_by_id(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM stations WHERE id=?", (station_id,))
        row = cur.fetchone()
        cols = [desc[0] for desc in cur.description] if cur.description else []
        conn.close()
        if row:
            return dict(zip(cols, row))
        return None
    except Exception as e:
        st.error(f"Erro ao carregar estação: {e}")
        return None


def get_all_stations():
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM stations ORDER BY id DESC")
        rows = cur.fetchall()
        cols = [desc[0] for desc in cur.description] if cur.description else []
        conn.close()
        return [dict(zip(cols, row)) for row in rows]
    except Exception as e:
        st.error(f"Erro ao listar estações: {e}")
        return []


def get_structures(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM structures WHERE station_id=?", (station_id,))
        rows = cur.fetchall()
        cols = [desc[0] for desc in cur.description] if cur.description else []
        conn.close()
        return [dict(zip(cols, row)) for row in rows]
    except Exception as e:
        st.error(f"Erro ao listar estruturas: {e}")
        return []


def get_samples(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM samples WHERE station_id=?", (station_id,))
        rows = cur.fetchall()
        cols = [desc[0] for desc in cur.description] if cur.description else []
        conn.close()
        return [dict(zip(cols, row)) for row in rows]
    except Exception as e:
        st.error(f"Erro ao listar amostras: {e}")
        return []


def get_photos(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM photos WHERE station_id=?", (station_id,))
        rows = cur.fetchall()
        cols = [desc[0] for desc in cur.description] if cur.description else []
        conn.close()
        return [dict(zip(cols, row)) for row in rows]
    except Exception as e:
        st.error(f"Erro ao listar fotos: {e}")
        return []


def insert_structures(station_id, structures):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        for s in structures:
            cur.execute(
                "INSERT INTO structures (station_id, tipo, strike, dip, dip_dir, plunge, azimuth, observacoes, created_at) VALUES (?,?,?,?,?,?,?,?,?)",
                (station_id, s.get("tipo"), s.get("strike"), s.get("dip"), s.get("dip_dir"), s.get("plunge"), s.get("azimuth"), s.get("observacoes"), now),
            )
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao salvar estruturas: {e}")


def insert_samples(station_id, samples):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        for s in samples:
            cur.execute(
                "INSERT INTO samples (station_id, codigo, tipo, finalidade, orientada, observacoes, created_at) VALUES (?,?,?,?,?,?,?)",
                (station_id, s.get("codigo"), s.get("tipo"), s.get("finalidade"), int(s.get("orientada") or 0), s.get("observacoes"), now),
            )
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao salvar amostras: {e}")


def insert_photos(station_id, photos):
    try:
        conn = get_connection()
        cur = conn.cursor()
        now = datetime.datetime.now().isoformat()
        for p in photos:
            cur.execute(
                "INSERT INTO photos (station_id, descricao, arquivo, created_at) VALUES (?,?,?,?)",
                (station_id, p.get("descricao"), p.get("arquivo"), now),
            )
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao salvar fotos: {e}")


def delete_structures_by_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM structures WHERE station_id=?", (station_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao excluir estruturas: {e}")


def delete_samples_by_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM samples WHERE station_id=?", (station_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao excluir amostras: {e}")


def delete_photos_by_station(station_id):
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM photos WHERE station_id=?", (station_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Erro ao excluir fotos: {e}")


# ---------------------------------------------------------------------------
# EXPORTAÇÃO
# ---------------------------------------------------------------------------
def to_csv_bytes(stations):
    try:
        if not stations:
            return b""
        df = pd.DataFrame(stations)
        for col in ["utm_east", "utm_north", "latitude", "longitude", "altitude"]:
            if col not in df.columns:
                df[col] = None
        out = io.StringIO()
        df.to_csv(out, index=False)
        return out.getvalue().encode("utf-8")
    except Exception as e:
        st.error(f"Erro ao exportar CSV: {e}")
        return b""


def to_geojson_bytes(stations):
    try:
        features = []
        for s in stations:
            lat = s.get("latitude")
            lon = s.get("longitude")
            if lat is None or lon is None:
                continue
            props = {k: v for k, v in s.items() if k not in ("latitude", "longitude")}
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [lon, lat]},
                "properties": props,
            })
        geojson = {"type": "FeatureCollection", "features": features}
        return json.dumps(geojson, ensure_ascii=False, indent=2).encode("utf-8")
    except Exception as e:
        st.error(f"Erro ao exportar GeoJSON: {e}")
        return b""


def escape_xml(value):
    if value is None:
        return ""
    value = str(value)
    value = value.replace("&", "&amp;")
    value = value.replace("<", "&lt;")
    value = value.replace(">", "&gt;")
    value = value.replace('"', "&quot;")
    value = value.replace("'", "&apos;")
    return value


def to_kml_bytes(stations):
    try:
        lines = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<kml xmlns="http://www.opengis.net/kml/2.2">',
            '<Document>',
            '<name>AfloraGeo - Estações</name>',
        ]
        for s in stations:
            lat = s.get("latitude")
            lon = s.get("longitude")
            if lat is None or lon is None:
                continue
            name = s.get("ponto_id") or "Estação"
            desc = f"{s.get('localizacao') or ''} | {s.get('litologia_principal') or ''} | {s.get('tipo_afloramento') or ''}"
            lines.append("<Placemark>")
            lines.append(f"<name>{escape_xml(name)}</name>")
            lines.append(f"<description>{escape_xml(desc)}</description>")
            lines.append("<Point>")
            alt = s.get('altitude') or 0
            lines.append(f"<coordinates>{lon},{lat},{alt}</coordinates>")
            lines.append("</Point>")
            lines.append("</Placemark>")
        lines += ["</Document>", "</kml>"]
        return "\n".join(lines).encode("utf-8")
    except Exception as e:
        st.error(f"Erro ao exportar KML: {e}")
        return b""


def to_docx_bytes(stations):
    try:
        doc = Document()
        title = doc.add_heading("AfloraGeo - Caderneta de Campo Geológica", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Relatório gerado em {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()
        for s in stations:
            doc.add_heading(s.get("ponto_id") or "Estação", level=2)
            p = doc.add_paragraph()
            p.add_run("Data: ").bold = True
            p.add_run(str(s.get("data") or "") + "\n")
            p.add_run("Localização: ").bold = True
            p.add_run(str(s.get("localizacao") or "") + "\n")
            p.add_run("Município: ").bold = True
            p.add_run(str(s.get("municipio") or "") + "\n")
            p.add_run("Tipo de afloramento: ").bold = True
            p.add_run(str(s.get("tipo_afloramento") or "") + "\n")
            p.add_run("Litologia principal: ").bold = True
            p.add_run(str(s.get("litologia_principal") or "") + "\n")
            p.add_run("Litologia secundária: ").bold = True
            p.add_run(str(s.get("litologia_secundaria") or "") + "\n")
            p.add_run("Granulometria: ").bold = True
            p.add_run(str(s.get("granulometria") or "") + "\n")
            p.add_run("Cor: ").bold = True
            p.add_run(str(s.get("cor") or "") + "\n")
            p.add_run("Intemperismo: ").bold = True
            p.add_run(str(s.get("intemperismo") or "") + "\n")
            p.add_run("Coordenadas UTM: ").bold = True
            p.add_run(f"Zona {s.get('utm_zone') or ''} {s.get('hemisferio') or ''} | E {s.get('utm_east') or ''} | N {s.get('utm_north') or ''}\n")
            p.add_run("Latitude/Longitude: ").bold = True
            p.add_run(f"{s.get('latitude') or ''}, {s.get('longitude') or ''}\n")
            p.add_run("Altitude: ").bold = True
            p.add_run(str(s.get("altitude") or "") + "\n")
            p.add_run("Acesso: ").bold = True
            p.add_run(str(s.get("acesso") or "") + "\n")
            p.add_run("Contexto geológico: ").bold = True
            p.add_run(str(s.get("contexto_geologico") or "") + "\n")
            p.add_run("Dimensões: ").bold = True
            p.add_run(str(s.get("dimensoes") or "") + "\n")
            p.add_run("Orientação do afloramento: ").bold = True
            p.add_run(str(s.get("orientacao_afloramento") or "") + "\n")
            p.add_run("Observações: ").bold = True
            p.add_run(str(s.get("observacoes") or "") + "\n")
            structs = get_structures(s.get("id"))
            if structs:
                doc.add_paragraph().add_run("Estruturas:").bold = True
                for es in structs:
                    doc.add_paragraph(
                        f"- {es.get('tipo') or ''}: strike {es.get('strike') or ''}, dip {es.get('dip') or ''}, "
                        f"dip_dir {es.get('dip_dir') or ''}, plunge {es.get('plunge') or ''}, azimuth {es.get('azimuth') or ''}. "
                        f"{es.get('observacoes') or ''}", style="List Bullet"
                    )
            samps = get_samples(s.get("id"))
            if samps:
                doc.add_paragraph().add_run("Amostras:").bold = True
                for sa in samps:
                    doc.add_paragraph(
                        f"- {sa.get('codigo') or ''} | {sa.get('tipo') or ''} | {sa.get('finalidade') or ''} | "
                        f"Orientada: {'Sim' if sa.get('orientada') else 'Não'}. {sa.get('observacoes') or ''}", style="List Bullet"
                    )
            doc.add_paragraph()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        st.error(f"Erro ao exportar DOCX: {e}")
        return b""


# ---------------------------------------------------------------------------
# UI COMPONENTES
# ---------------------------------------------------------------------------
def sidebar():
    st.sidebar.markdown("# AfloraGeo 🗻")
    st.sidebar.markdown("Caderneta de Campo Geológica")
    st.sidebar.markdown("---")
    if "pagina" not in st.session_state:
        st.session_state.pagina = "Nova estação"
    pages = [
        "Nova estação", "Lista de estações", "Mapa",
        "Tabelas de apoio", "Exportar", "Premium", "Configurações",
    ]
    current_idx = pages.index(st.session_state.pagina) if st.session_state.pagina in pages else 0
    selected = st.sidebar.radio("Navegação", pages, index=current_idx)
    if selected != st.session_state.pagina:
        st.session_state.pagina = selected
        if "edit_id" in st.session_state:
            st.session_state.edit_id = None
        safe_rerun()
    st.sidebar.markdown("---")
    lic = get_license()
    count = get_station_count()
    if lic and lic.get("is_premium") == 1:
        st.sidebar.markdown("<span class='premium-badge'>⭐ PREMIUM ATIVADO</span>", unsafe_allow_html=True)
        st.sidebar.markdown(f"Plano: {lic.get('plan_type') or ''}")
        st.sidebar.markdown(f"Expira em: {lic.get('expires_at') or ''}")
    else:
        limit = lic.get("stations_limit", DEFAULT_FREE_LIMIT) if lic else DEFAULT_FREE_LIMIT
        st.sidebar.markdown("<span class='free-badge'>Versão Gratuita</span>", unsafe_allow_html=True)
        st.sidebar.markdown(f"{count} / {limit} estações")
        progress = min(count / limit, 1.0) if limit and limit > 0 else 1.0
        st.sidebar.progress(progress)


# ---------------------------------------------------------------------------
# PÁGINA: NOVA ESTAÇÃO
# ---------------------------------------------------------------------------
def pagina_nova_estacao():
    st.title("Nova Estação 📝")
    if not can_add_station():
        st.warning("Você atingiu o limite de estações gratuitas. Ative o Premium para continuar.")
        return
    edit_id = st.session_state.get("edit_id")
    station = None
    if edit_id:
        station = get_station_by_id(edit_id)
        if station:
            st.info(f"Editando estação {station.get('ponto_id') or edit_id}")
        else:
            st.session_state.edit_id = None

    # Conversor UTM → Lat/Lon fora do formulário
    with st.expander("Conversão UTM → Lat/Lon", expanded=False):
        c_utm1, c_utm2, c_utm3 = st.columns(3)
        with c_utm1:
            conv_east = st.number_input("UTM East (E)", value=0.0, key="conv_utm_east")
        with c_utm2:
            conv_north = st.number_input("UTM North (N)", value=0.0, key="conv_utm_north")
        with c_utm3:
            conv_zone = st.text_input("Zona UTM", value="23", key="conv_utm_zone")
        conv_hemi = st.selectbox("Hemisfério", ["S", "N"], key="conv_utm_hemi")

        if st.button("Converter UTM → Lat/Lon", key="btn_converter"):
            lat_c, lon_c = utm_to_latlon(conv_zone, conv_east, conv_north, conv_hemi)
            if lat_c is not None and lon_c is not None:
                st.session_state["lat_conv"] = lat_c
                st.session_state["lon_conv"] = lon_c
                st.success(f"Lat: {lat_c:.6f}, Lon: {lon_c:.6f}")
            else:
                st.error("Conversão inválida. Verifique zona e coordenadas.")

    with st.form("station_form"):
        col1, col2 = st.columns(2)
        with col1:
            ponto_id = st.text_input("Ponto ID", value=station.get("ponto_id") if station else next_station_id())
        with col2:
            data_field = station.get("data") if station else datetime.date.today().isoformat()
            data = st.date_input("Data", value=datetime.date.fromisoformat(data_field) if data_field else datetime.date.today())
        st.subheader("Localização")
        c1, c2, c3 = st.columns(3)
        with c1:
            utm_zone = st.text_input("Zona UTM", value=station.get("utm_zone") if station else "23")
        with c2:
            hemisferio = st.selectbox("Hemisfério", ["S", "N"], index=0 if (not station or station.get("hemisferio") == "S") else 1)
        with c3:
            altitude = st.number_input("Altitude (m)", value=float(station.get("altitude") or 0.0), step=1.0)
        c4, c5 = st.columns(2)
        with c4:
            utm_east = st.number_input("Coordenada Leste (E)", value=float(station.get("utm_east") or 0.0), step=0.1)
        with c5:
            utm_north = st.number_input("Coordenada Norte (N)", value=float(station.get("utm_north") or 0.0), step=0.1)
        lat = st.number_input("Latitude", value=float(station.get("latitude") if station else st.session_state.get("lat_conv", 0.0)), format="%.6f", step=0.000001)
        lon = st.number_input("Longitude", value=float(station.get("longitude") if station else st.session_state.get("lon_conv", 0.0)), format="%.6f", step=0.000001)
        localizacao = st.text_input("Localização / Descrição do local", value=station.get("localizacao") or "")
        municipio = st.text_input("Município", value=station.get("municipio") or "")
        st.subheader("Caracterização do Afloramento")
        tipo_afloramento = st.selectbox("Tipo de afloramento", TIPOS_AFLO, index=TIPOS_AFLO.index(station.get("tipo_afloramento")) if station and station.get("tipo_afloramento") in TIPOS_AFLO else 0)
        contexto_geologico = st.text_area("Contexto geológico", value=station.get("contexto_geologico") or "")
        dimensoes = st.text_input("Dimensões (ex: 10m x 3m)", value=station.get("dimensoes") or "")
        orientacao_afloramento = st.text_input("Orientação do afloramento", value=station.get("orientacao_afloramento") or "")
        acesso_vals = station.get("acesso").split(", ") if station and station.get("acesso") else []
        acesso = st.multiselect("Meio de acesso", ACESSOS, default=[a for a in acesso_vals if a in ACESSOS])
        st.subheader("Litologia")
        litologia_principal = st.text_input("Litologia principal", value=station.get("litologia_principal") or "")
        litologia_secundaria = st.text_input("Litologia secundária", value=station.get("litologia_secundaria") or "")
        granulometria = st.text_input("Granulometria / Textura", value=station.get("granulometria") or "")
        cor_index = 0
        if station and station.get("cor"):
            stored = station.get("cor")
            for i, (cod, nome) in enumerate(CORES_MUNSELL):
                label = f"{cod} ({nome})" if cod else nome
                if label == stored:
                    cor_index = i
                    break
            else:
                # Fallback: match parcial por código ou nome
                for i, (cod, nome) in enumerate(CORES_MUNSELL):
                    if cod and (cod in stored or stored in cod):
                        cor_index = i
                        break
        cor_sel = st.selectbox("Cor (Munsell)", CORES_MUNSELL, index=cor_index, format_func=lambda x: f"{x[0]} - {x[1]}" if x[0] else x[1])
        cor = f"{cor_sel[0]} ({cor_sel[1]})" if cor_sel[0] else ""
        intemperismo_index = 0
        if station and station.get("intemperismo"):
            for i, (cod, _, _) in enumerate(INTEMPERISMO_IBGE):
                if station.get("intemperismo") == cod:
                    intemperismo_index = i
                    break
        intemperismo = st.selectbox("Intemperismo (IBGE-CPRM)", [i[0] for i in INTEMPERISMO_IBGE], index=intemperismo_index)
        observacoes = st.text_area("Observações gerais", value=station.get("observacoes") or "")

        st.subheader("Estruturas geológicas")
        structures = []
        with st.expander("Adicionar estruturas"):
            existing_structs = get_structures(edit_id) if edit_id else []
            n_structs = st.number_input("Quantidade de estruturas", min_value=0, max_value=20, value=len(existing_structs), step=1)
            for i in range(int(n_structs)):
                st.markdown(f"**Estrutura {i+1}**")
                es = existing_structs[i] if i < len(existing_structs) else {}
                tipo = st.selectbox("Tipo", TIPOS_ESTRUTURA_SELECT, index=TIPOS_ESTRUTURA_SELECT.index(es.get("tipo")) if es.get("tipo") in TIPOS_ESTRUTURA_SELECT else 0, key=f"st_tipo_{i}")
                cc1, cc2, cc3, cc4, cc5 = st.columns(5)
                with cc1:
                    strike = st.text_input("Strike", value=es.get("strike") or "", key=f"st_strike_{i}")
                with cc2:
                    dip = st.number_input("Dip", value=float(es.get("dip") or 0.0), step=1.0, key=f"st_dip_{i}")
                with cc3:
                    dip_dir = st.number_input("Dip dir", value=float(es.get("dip_dir") or 0.0), step=1.0, key=f"st_dipdir_{i}")
                with cc4:
                    plunge = st.number_input("Plunge", value=float(es.get("plunge") or 0.0), step=1.0, key=f"st_plunge_{i}")
                with cc5:
                    azimuth = st.number_input("Azimuth", value=float(es.get("azimuth") or 0.0), step=1.0, key=f"st_azimuth_{i}")
                obs_struct = st.text_input("Observações da estrutura", value=es.get("observacoes") or "", key=f"st_obs_{i}")
                structures.append({
                    "tipo": tipo, "strike": strike, "dip": dip, "dip_dir": dip_dir,
                    "plunge": plunge, "azimuth": azimuth, "observacoes": obs_struct,
                })

        st.subheader("Amostras")
        samples = []
        with st.expander("Adicionar amostras"):
            existing_samples = get_samples(edit_id) if edit_id else []
            n_samples = st.number_input("Quantidade de amostras", min_value=0, max_value=20, value=len(existing_samples), step=1)
            for i in range(int(n_samples)):
                st.markdown(f"**Amostra {i+1}**")
                sa = existing_samples[i] if i < len(existing_samples) else {}
                cod = st.text_input("Código", value=sa.get("codigo") or f"{ponto_id}-A{i+1}", key=f"sa_cod_{i}")
                tipo = st.text_input("Tipo de amostra", value=sa.get("tipo") or "Rocha", key=f"sa_tipo_{i}")
                finalidade = st.selectbox("Finalidade", FINALIDADES_AMOSTRA_LISTA, index=FINALIDADES_AMOSTRA_LISTA.index(sa.get("finalidade")) if sa.get("finalidade") in FINALIDADES_AMOSTRA_LISTA else 0, key=f"sa_final_{i}")
                orientada = st.checkbox("Orientada", value=bool(sa.get("orientada")), key=f"sa_orient_{i}")
                obs_samp = st.text_input("Observações da amostra", value=sa.get("observacoes") or "", key=f"sa_obs_{i}")
                samples.append({
                    "codigo": cod, "tipo": tipo, "finalidade": finalidade,
                    "orientada": orientada, "observacoes": obs_samp,
                })

        st.subheader("Fotos")
        photos = []
        with st.expander("Anexar fotos"):
            uploaded = st.file_uploader("Selecionar fotos", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
            if uploaded:
                os.makedirs("fotos", exist_ok=True)
                for uf in uploaded:
                    try:
                        filename = f"fotos/{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{uf.name}"
                        with open(filename, "wb") as f:
                            f.write(uf.read())
                        photos.append({"descricao": uf.name, "arquivo": filename})
                    except Exception as e:
                        st.error(f"Erro ao salvar foto: {e}")

        submitted = st.form_submit_button("Salvar estação")
        if submitted:
            try:
                data_str = data.isoformat() if hasattr(data, "isoformat") else str(data)
                station_data = {
                    "ponto_id": ponto_id,
                    "data": data_str,
                    "utm_zone": utm_zone,
                    "hemisferio": hemisferio,
                    "utm_east": utm_east,
                    "utm_north": utm_north,
                    "latitude": lat,
                    "longitude": lon,
                    "altitude": altitude,
                    "localizacao": localizacao,
                    "municipio": municipio,
                    "contexto_geologico": contexto_geologico,
                    "tipo_afloramento": tipo_afloramento,
                    "dimensoes": dimensoes,
                    "orientacao_afloramento": orientacao_afloramento,
                    "acesso": ", ".join(acesso),
                    "litologia_principal": litologia_principal,
                    "litologia_secundaria": litologia_secundaria,
                    "granulometria": granulometria,
                    "cor": cor,
                    "intemperismo": intemperismo,
                    "observacoes": observacoes,
                }
                if edit_id:
                    update_station(edit_id, station_data)
                    delete_structures_by_station(edit_id)
                    delete_samples_by_station(edit_id)
                    delete_photos_by_station(edit_id)
                    station_id = edit_id
                else:
                    station_id = insert_station(station_data)
                if station_id:
                    if structures:
                        insert_structures(station_id, structures)
                    if samples:
                        insert_samples(station_id, samples)
                    if photos:
                        insert_photos(station_id, photos)
                    st.success("Estação salva com sucesso!")
                    st.session_state.edit_id = None
                    st.balloons()
            except Exception as e:
                st.error(f"Erro ao salvar: {e}")


# ---------------------------------------------------------------------------
# PÁGINA: LISTA DE ESTAÇÕES
# ---------------------------------------------------------------------------
def pagina_lista_estacoes():
    st.title("Lista de Estações 📋")
    stations = get_all_stations()
    if not stations:
        st.info("Nenhuma estação cadastrada.")
        return
    df = pd.DataFrame(stations)
    col1, col2, col3 = st.columns(3)
    with col1:
        busca = st.text_input("Buscar texto")
    with col2:
        datas = sorted([d for d in df["data"].dropna().unique() if d])
        data_sel = st.selectbox("Filtrar por data", ["Todas"] + datas)
    with col3:
        litos = sorted([l for l in df["litologia_principal"].dropna().unique() if l])
        lito_sel = st.selectbox("Filtrar por litologia", ["Todas"] + litos)
    filtered = df
    if busca:
        mask = filtered.astype(str).apply(lambda row: busca.lower() in " ".join(row).lower(), axis=1)
        filtered = filtered[mask]
    if data_sel != "Todas":
        filtered = filtered[filtered["data"] == data_sel]
    if lito_sel != "Todas":
        filtered = filtered[filtered["litologia_principal"] == lito_sel]

    st.dataframe(
        filtered,
        column_config={
            "id": None,
            "ponto_id": "Ponto",
            "data": "Data",
            "localizacao": "Localização",
            "municipio": "Município",
            "tipo_afloramento": "Tipo",
            "litologia_principal": "Litologia",
            "litologia_secundaria": "Lit. Secundária",
            "cor": "Cor",
            "utm_east": st.column_config.NumberColumn("E (UTM)", format="%.1f"),
            "utm_north": st.column_config.NumberColumn("N (UTM)", format="%.1f"),
            "latitude": st.column_config.NumberColumn("Lat", format="%.6f"),
            "longitude": st.column_config.NumberColumn("Lon", format="%.6f"),
            "altitude": st.column_config.NumberColumn("Alt (m)", format="%.1f"),
            "created_at": None,
            "updated_at": None,
        },
        use_container_width=True,
        hide_index=True,
    )

    for _, row in filtered.iterrows():
        sid = row.get("id")
        ponto = row.get("ponto_id") or f"ID {sid}"
        with st.expander(f"{ponto} | {row.get('data') or ''} | {row.get('localizacao') or ''}"):
            st.write(f"**Município:** {row.get('municipio') or ''}")
            st.write(f"**Tipo:** {row.get('tipo_afloramento') or ''}")
            st.write(f"**Litologia:** {row.get('litologia_principal') or ''}")
            st.write(f"**Cor:** {row.get('cor') or ''}")
            st.write(f"**UTM:** Zona {row.get('utm_zone') or ''} {row.get('hemisferio') or ''} E {row.get('utm_east') or ''} N {row.get('utm_north') or ''}")
            st.write(f"**Lat/Lon:** {row.get('latitude') or ''}, {row.get('longitude') or ''}")
            st.write(f"**Observações:** {row.get('observacoes') or ''}")
            structs = get_structures(sid)
            if structs:
                st.write("**Estruturas:**")
                for es in structs:
                    st.write(f"- {es.get('tipo')}: strike {es.get('strike')}, dip {es.get('dip')}")
            samps = get_samples(sid)
            if samps:
                st.write("**Amostras:**")
                for sa in samps:
                    st.write(f"- {sa.get('codigo')} | {sa.get('finalidade')} | Orientada: {'Sim' if sa.get('orientada') else 'Não'}")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Editar", key=f"edit_{sid}"):
                    st.session_state.edit_id = sid
                    st.session_state.pagina = "Nova estação"
                    safe_rerun()
            with c2:
                if st.button("Excluir", key=f"del_{sid}"):
                    if delete_station(sid):
                        st.success("Estação excluída.")
                        safe_rerun()


# ---------------------------------------------------------------------------
# PÁGINA: MAPA
# ---------------------------------------------------------------------------
def pagina_mapa():
    st.title("Mapa 🗺️")
    stations = get_all_stations()
    valid = [s for s in stations if s.get("latitude") and s.get("longitude")]
    if not valid:
        st.info("Nenhuma estação com coordenadas válidas.")
        return
    m = folium.Map(location=[-14.235, -51.925], zoom_start=4, tiles="OpenStreetMap")
    marker_cluster = MarkerCluster().add_to(m)
    for s in valid:
        popup_html = f"""
        <b>{s.get('ponto_id') or 'Estação'}</b><br>
        {s.get('localizacao') or ''}<br>
        {s.get('tipo_afloramento') or ''}<br>
        {s.get('litologia_principal') or ''}
        """
        folium.Marker(
            [s.get("latitude"), s.get("longitude")],
            popup=folium.Popup(popup_html, max_width=250),
            tooltip=s.get("ponto_id") or "Estação",
        ).add_to(marker_cluster)
    st_folium(m, width=1200, height=700)


# ---------------------------------------------------------------------------
# PÁGINA: TABELAS DE APOIO
# ---------------------------------------------------------------------------
def pagina_tabelas():
    st.title("Tabelas de Apoio 📊")
    tabs = st.tabs(["Densidade", "Vp", "Susceptibilidade", "Declividade", "Drenagem", "Relevo", "Intemperismo"])
    with tabs[0]:
        st.subheader("Densidade de rochas (Telford 1990, Kearey 2002)")
        st.dataframe(pd.DataFrame(DENSIDADE_ROCHAS, columns=["Rocha", "Intervalo (g/cm³)", "Média (g/cm³)"]), use_container_width=True)
    with tabs[1]:
        st.subheader("Velocidade de Ondas P (km/s) - Kearey, Brooks & Hill, 2002")
        st.dataframe(pd.DataFrame(VELOCIDADE_ONDAS_P, columns=["Material", "Vp (km/s)"]), use_container_width=True, hide_index=True)
    with tabs[2]:
        st.subheader("Suscetibilidade Magnética (×10⁻³ SI) - CPRM")
        st.dataframe(pd.DataFrame(SUSCETIBILIDADE_MAGNETICA, columns=["Material", "Susc. (×10⁻³ SI)"]), use_container_width=True, hide_index=True)
    with tabs[3]:
        st.subheader("Classes de declividade (IBGE)")
        st.dataframe(pd.DataFrame(CLASSES_DECLIVIDADE, columns=["Classe", "Declividade", "Significado"]), use_container_width=True)
    with tabs[4]:
        st.subheader("Padrões de drenagem")
        st.dataframe(pd.DataFrame(PADROES_DRENAGEM, columns=["Padrão", "Característica", "Contexto"]), use_container_width=True)
    with tabs[5]:
        st.subheader("Formas de relevo (IBGE)")
        st.dataframe(pd.DataFrame(FORMAS_RELEVO, columns=["Forma", "Sigla", "Descrição"]), use_container_width=True)
    with tabs[6]:
        st.subheader("Intemperismo (IBGE-CPRM)")
        st.dataframe(pd.DataFrame(INTEMPERISMO_IBGE, columns=["Código", "Descrição", "Observação"]), use_container_width=True)


# ---------------------------------------------------------------------------
# PÁGINA: EXPORTAR
# ---------------------------------------------------------------------------
def pagina_exportar():
    st.title("Exportar 📤")
    stations = get_all_stations()
    if not stations:
        st.info("Nenhuma estação para exportar.")
        return
    st.write(f"Total de estações: {len(stations)}")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.download_button("Baixar CSV", data=to_csv_bytes(stations), file_name="aflorageo_estacoes.csv", mime="text/csv")
    with col2:
        st.download_button("Baixar GeoJSON", data=to_geojson_bytes(stations), file_name="aflorageo_estacoes.geojson", mime="application/geo+json")
    with col3:
        st.download_button("Baixar KML", data=to_kml_bytes(stations), file_name="aflorageo_estacoes.kml", mime="application/vnd.google-earth.kml+xml")
    with col4:
        if is_premium():
            st.download_button("Baixar DOCX", data=to_docx_bytes(stations), file_name="aflorageo_relatorio.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.button("DOCX (Premium)", disabled=True)
            st.caption("DOCX disponível apenas para usuários Premium.")


# ---------------------------------------------------------------------------
# PÁGINA: PREMIUM
# ---------------------------------------------------------------------------
def pagina_premium():
    st.title("Premium 💎")

    if "plano_selecionado" not in st.session_state:
        st.session_state.plano_selecionado = "Anual"

    lic = get_license()
    if lic and lic.get("is_premium") == 1:
        st.success(f"Você já é Premium! Plano: {lic.get('plan_type')}. Expira em {lic.get('expires_at')}.")
        return

    st.markdown("Escolha um plano para desbloquear exportação DOCX e estações ilimitadas.")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("### Mensal")
        st.markdown("R$ 19,90 / mês")
        if st.button("Assinar Mensal"):
            st.session_state.plano_selecionado = "Mensal"
            safe_rerun()
    with c2:
        st.markdown("### Semestral")
        st.markdown("R$ 99,00 / 6 meses")
        if st.button("Assinar Semestral"):
            st.session_state.plano_selecionado = "Semestral"
            safe_rerun()
    with c3:
        st.markdown("### Anual")
        st.markdown("R$ 179,00 / ano")
        if st.button("Assinar Anual"):
            st.session_state.plano_selecionado = "Anual"
            safe_rerun()

    with st.form("activation_form"):
        st.markdown("Ou ative com chave de licença:")
        chave = st.text_input("Chave de ativação", type="password")
        email = st.text_input("E-mail")
        submitted = st.form_submit_button("Ativar")
        if submitted:
            if chave.strip() == ACTIVATION_KEY:
                plano_ativo = st.session_state.get("plano_selecionado", "Anual")
                if activate_premium(plano_ativo, email):
                    st.success(f"Premium ativado com sucesso! Plano: {plano_ativo}")
                    safe_rerun()
            else:
                st.error("Chave de ativação inválida.")
    st.markdown("---")
    st.markdown("### Comparação de planos")
    comp = pd.DataFrame({
        "Recurso": ["Estações", "Exportar CSV/GeoJSON/KML", "Exportar DOCX", "Suporte"],
        "Gratuito": ["30", "Sim", "Não", "Comunitário"],
        "Premium": ["Ilimitado", "Sim", "Sim", "E-mail"],
    })
    st.dataframe(comp, use_container_width=True, hide_index=True)


# ---------------------------------------------------------------------------
# PÁGINA: CONFIGURAÇÕES
# ---------------------------------------------------------------------------
def pagina_configuracoes():
    st.title("Configurações ⚙️")
    lic = get_license()
    st.subheader("Status da licença")
    if lic and lic.get("is_premium") == 1:
        st.markdown("<span class='premium-badge'>⭐ Premium ativo</span>", unsafe_allow_html=True)
        st.write(f"Plano: {lic.get('plan_type')}")
        st.write(f"E-mail: {lic.get('user_email') or 'Não informado'}")
        st.write(f"Expira em: {lic.get('expires_at')}")
    else:
        st.markdown("<span class='free-badge'>Versão gratuita</span>", unsafe_allow_html=True)
        st.write(f"Limite: {lic.get('stations_limit', DEFAULT_FREE_LIMIT) if lic else DEFAULT_FREE_LIMIT} estações")
    st.subheader("Conta")
    st.write("Banco de dados local: " + DB_PATH)
    st.subheader("Zona de perigo")
    with st.expander("Redefinir dados"):
        st.warning("Atenção: esta ação excluirá todas as estações e redefinirá a licença para gratuita.")
        if st.button("Resetar todos os dados"):
            try:
                conn = get_connection()
                cur = conn.cursor()
                for t in ["structures", "samples", "photos", "stations"]:
                    cur.execute(f"DELETE FROM {t}")
                conn.commit()
                conn.close()
                reset_license()
                st.success("Dados redefinidos.")
                safe_rerun()
            except Exception as e:
                st.error(f"Erro ao resetar dados: {e}")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    st.set_page_config(page_title="AfloraGeo", page_icon="🗻", layout="wide")
    inject_css()
    init_db()
    sidebar()
    pagina = st.session_state.get("pagina", "Nova estação")
    if pagina == "Nova estação":
        pagina_nova_estacao()
    elif pagina == "Lista de estações":
        pagina_lista_estacoes()
    elif pagina == "Mapa":
        pagina_mapa()
    elif pagina == "Tabelas de apoio":
        pagina_tabelas()
    elif pagina == "Exportar":
        pagina_exportar()
    elif pagina == "Premium":
        pagina_premium()
    elif pagina == "Configurações":
        pagina_configuracoes()
    else:
        pagina_nova_estacao()


if __name__ == "__main__":
    main()